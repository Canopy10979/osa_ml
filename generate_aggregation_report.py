from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# PATHS
# ============================================================

ROOT = Path.cwd()

MENTOR_DIR = ROOT / "mentor_action_items"
MENTOR_DIR.mkdir(parents=True, exist_ok=True)

AGG_CODE = ROOT / "aggregation_code.txt"
MODEL_CODE = ROOT / "aggregation_model_code.txt"
RESULTS = ROOT / "subject_level_results.txt"
WINDOWS = ROOT / "window_predictions_sample.txt"
SEVERITY = ROOT / "severity_logic.txt"


# ============================================================
# IMPORTANT:
# Always create a NEW filename.
# This avoids PermissionError if an older Word file is open.
# ============================================================

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

OUTPUT = MENTOR_DIR / (
    f"02_patient_level_aggregation_{timestamp}.docx"
)


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def read_text_file(path):
    """
    Tries common encodings so PowerShell-generated text files
    can still be read correctly.
    """

    if not path.exists():
        return (
            f"[File not found: {path.name}]\n"
            "This section could not be populated."
        )

    encodings = [
        "utf-8",
        "utf-16",
        "utf-8-sig",
        "cp1252"
    ]

    for encoding in encodings:
        try:
            return path.read_text(
                encoding=encoding
            )
        except UnicodeError:
            pass

    return path.read_text(
        errors="ignore"
    )


def add_code_block(document, text):
    """
    Adds code/output using a monospace font.
    """

    paragraph = document.add_paragraph()

    run = paragraph.add_run(text)

    run.font.name = "Consolas"
    run.font.size = Pt(8)


def add_bullet(document, text):
    document.add_paragraph(
        text,
        style="List Bullet"
    )


# ============================================================
# CREATE DOCUMENT
# ============================================================

document = Document()


# -----------------------------
# Page formatting
# -----------------------------

section = document.sections[0]

section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)


# -----------------------------
# Default font
# -----------------------------

normal_style = document.styles["Normal"]

normal_style.font.name = "Arial"
normal_style.font.size = Pt(10.5)


# ============================================================
# TITLE
# ============================================================

title = document.add_heading(
    "Patient-Level Aggregation Findings",
    level=0
)

title.alignment = WD_ALIGN_PARAGRAPH.CENTER


subtitle = document.add_paragraph(
    "OSA Machine Learning Project\n"
    "STEM Mentoring Action Item #2"
)

subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER


document.add_paragraph(
    "This report documents how the current OSA machine-learning "
    "codebase converts individual time-window predictions into "
    "subject-level outputs and investigates whether those outputs "
    "are further translated into severity scores."
)


# ============================================================
# SECTION 1
# ============================================================

document.add_heading(
    "1. Window-Level Predictions",
    level=1
)

document.add_paragraph(
    "The machine-learning pipeline first generates predictions "
    "for individual time windows or epochs. These window-level "
    "outputs serve as the basic units that are later grouped by "
    "subject."
)

add_bullet(
    document,
    "Each prediction corresponds to an individual time segment."
)

add_bullet(
    document,
    "Predictions can include a binary predicted class and/or a "
    "predicted probability."
)

add_bullet(
    document,
    "Multiple window predictions belonging to the same subject "
    "are later aggregated."
)

document.add_paragraph(
    "Sample out-of-fold predictions:"
)

add_code_block(
    document,
    read_text_file(WINDOWS)
)


# ============================================================
# SECTION 2
# ============================================================

document.add_heading(
    "2. Subject-Level Aggregation Method",
    level=1
)

document.add_paragraph(
    "Codebase inspection identified explicit grouping of "
    "window-level results by subject. The relevant implementation "
    "is shown below."
)

add_code_block(
    document,
    read_text_file(AGG_CODE)
)

document.add_paragraph(
    "The presence of a groupby('subject') operation confirms that "
    "predictions from multiple windows belonging to the same "
    "individual are combined before subject-level results are "
    "written."
)


# ============================================================
# SECTION 3
# ============================================================

document.add_heading(
    "3. Aggregation Formula",
    level=1
)

document.add_paragraph(
    "Additional inspection of the modeling code identified "
    "subject-level aggregation using grouped averages."
)

add_code_block(
    document,
    read_text_file(MODEL_CODE)
)

document.add_paragraph(
    "Based on the inspected code, the pipeline includes subject-"
    "level calculations based on the mean of window-level values."
)

add_bullet(
    document,
    "Mean predicted output: average prediction across windows "
    "belonging to the same subject."
)

add_bullet(
    document,
    "Observed fraction: average of the true binary apnea labels "
    "across windows belonging to the same subject."
)

document.add_paragraph(
    "For a binary label, taking the mean of the labels is "
    "equivalent to calculating the fraction of windows labeled "
    "positive."
)


# ============================================================
# SECTION 4
# ============================================================

document.add_heading(
    "4. Severity Scoring Method",
    level=1
)

document.add_paragraph(
    "The codebase was searched for explicit severity-related "
    "logic including AHI, mild, moderate, severe, and other "
    "severity terminology."
)

add_code_block(
    document,
    read_text_file(SEVERITY)
)

document.add_paragraph(
    "A model-derived apnea fraction or mean predicted probability "
    "must not automatically be interpreted as a clinical "
    "Apnea-Hypopnea Index (AHI). Clinical severity categories "
    "should only be assigned when explicit code or ground-truth "
    "clinical labels support those thresholds."
)


# ============================================================
# SECTION 5
# ============================================================

document.add_heading(
    "5. Subject-Level Results",
    level=1
)

document.add_paragraph(
    "The current pipeline produces saved subject-level prediction "
    "results. These results provide the direct output of the "
    "aggregation process."
)

add_code_block(
    document,
    read_text_file(RESULTS)
)


# ============================================================
# SECTION 6
# ============================================================

document.add_heading(
    "6. Validation and Limitations",
    level=1
)

document.add_paragraph(
    "The aggregation method was validated by tracing the pipeline "
    "from window-level predictions through subject grouping and "
    "into the saved subject-level output."
)

add_bullet(
    document,
    "Window-level predictions exist before subject aggregation."
)

add_bullet(
    document,
    "The code explicitly groups records by subject."
)

add_bullet(
    document,
    "The grouped calculations include averages and counts of "
    "predicted or observed positive windows."
)

add_bullet(
    document,
    "Subject-level predictions are written to a dedicated "
    "subject_level_predictions.csv file."
)

document.add_paragraph(
    "Limitations:"
)

add_bullet(
    document,
    "A positive-window fraction is not automatically equivalent "
    "to clinically measured AHI."
)

add_bullet(
    document,
    "Mean predicted probability represents model confidence or "
    "risk rather than a direct clinical severity measurement."
)

add_bullet(
    document,
    "Subject-level performance should ideally be validated "
    "against independent clinical severity labels."
)


# ============================================================
# SECTION 7
# ============================================================

document.add_heading(
    "7. Conclusion",
    level=1
)

document.add_paragraph(
    "The current OSA machine-learning codebase does perform "
    "patient-level aggregation. Individual time-window "
    "predictions are grouped by subject, after which summary "
    "statistics such as predicted-positive counts, observed "
    "fractions, and mean prediction values are calculated."
)

document.add_paragraph(
    "This provides a pathway from time-window classification to "
    "subject-level assessment. However, subject-level prediction "
    "scores should be distinguished from clinical OSA severity "
    "unless they are explicitly calibrated against AHI or other "
    "clinical severity measurements."
)


# ============================================================
# SAVE
# ============================================================

try:

    document.save(OUTPUT)

except PermissionError:

    # Extremely unlikely now because filename includes timestamp,
    # but this gives us a second fallback anyway.

    fallback = MENTOR_DIR / (
        f"02_patient_level_aggregation_backup_{timestamp}.docx"
    )

    document.save(fallback)

    OUTPUT = fallback


print("\n✅ SUCCESS!")
print("Word report created at:")
print(OUTPUT.resolve())

print("\n📄 Open it with:")
print(
    f'Start-Process "{OUTPUT}"'
)
